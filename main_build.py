import os
import time
import shutil
from copy import copy
import yaml
from datetime import datetime, timedelta
from typing import Optional, List

from pydantic import BaseModel
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from openpyxl import load_workbook
from openpyxl import Workbook
from tkinter import filedialog, messagebox
import win32com.client

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import openpyxl
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from datetime import datetime
from selenium.webdriver.common.action_chains import ActionChains


class PathConfig(BaseModel):
    output: str
    chrome_driver: str
    download: str


class IplasConfig(BaseModel):
    username: str
    password: str


class EmailConfig(BaseModel):
    TO: List[str]
    CC: List[str]


class ReportTimeConfig(BaseModel):
    start_time: Optional[str]
    end_time: Optional[str]


class AppConfig(BaseModel):
    domains: List[str]
    path: PathConfig
    iplas: IplasConfig
    email: EmailConfig
    time: ReportTimeConfig


def load_yaml(file_path: str):
    with open(file_path, "r", encoding="utf-8") as file:
        return yaml.safe_load(file)


def load_app_config(path: str) -> AppConfig:
    config = load_yaml(path)
    domains = config.get("domains", [])

    _path = config["path_config"]
    path_config = PathConfig(
        output=_path["output"],
        chrome_driver=_path["chromedriver"],
        download=_path["download"]
    )

    _iplas = config["iplas_config"]
    iplas_config = IplasConfig(
        username=_iplas["username"],
        password=_iplas["password"]
    )

    _email = config["email_config"]
    email_config = EmailConfig(
        TO=_email["TO"],
        CC=_email["CC"]
    )

    _time = config["report_time"]
    report_time_config = ReportTimeConfig(
        start_time=_time["start_time"],
        end_time=_time["end_time"]
    )

    return AppConfig(
        domains=domains,
        path=path_config,
        iplas=iplas_config,
        email=email_config,
        time=report_time_config
    )


def rgb_to_hex(rgb):
    if hasattr(rgb, 'rgb'):
        return '#{:06x}'.format(rgb.rgb)
    elif hasattr(rgb, 'red'):
        return '#{:02x}{:02x}{:02x}'.format(rgb.red, rgb.green, rgb.blue)
    else:
        return '#000000'


def get_run_style(rPr):
    style = ''
    if rPr.xpath('.//w:b'):
        style += 'font-weight:bold;'
    if rPr.xpath('.//w:i'):
        style += 'font-style:italic;'
    if rPr.xpath('.//w:u'):
        style += 'text-decoration:underline;'
    return style


def process_paragraph(paragraph):
    p_html = "<p style='margin:0;'>"
    for run in paragraph.xpath('.//w:r'):
        text = ''.join([t.text for t in run.xpath('.//w:t')])
        rPr = run.xpath('.//w:rPr')
        if rPr:
            style = get_run_style(rPr[0])
        else:
            style = ''
        if any(project in text for project, _, _ in data):
            style += 'font-weight:bold;'
        p_html += f"<span style='{style}'>{text}</span>"
    return p_html + "</p>"


if __name__ == "__main__":
    config_path = "config.yaml"
    date_time_format = "%Y/%m/%d %H:%M:%S"
    base_url = "http://10.177.240.149/iPLAS/main"
    max_download_wait_time = 60  # Maximum wait time in seconds

    # Define the output folder path
    try:
        config = load_app_config(config_path)
        domains = config.domains
        output_folder = config.path.output

        # Check if the output folder exists, create it if it doesn't
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        else:
            # Remove all files in the output folder
            for file_name in os.listdir(output_folder):
                file_path = os.path.join(output_folder, file_name)
                if os.path.isfile(file_path):
                    os.remove(file_path)

        # Define the list of domains
        chrome_options = Options()
        # Set up the webdriver
        chrome_service = Service(executable_path=config.path.chrome_driver)
        for desired_url in domains:
            chrome_options.add_argument(f"--unsafely-treat-insecure-origin-as-secure={desired_url}")
        driver = webdriver.Chrome(service=chrome_service, options=chrome_options)

        # Navigate to the website and login
        driver.get(base_url)
        wait = WebDriverWait(driver, 10)
        wait.until(EC.presence_of_element_located((By.CLASS_NAME, 'pega-login-username')))
        username_input = driver.find_element(By.CLASS_NAME, 'pega-login-username')
        password_input = driver.find_element(By.CLASS_NAME, 'pega-login-password')
        username_input.send_keys(config.iplas.username)
        password_input.send_keys(config.iplas.password)
        login_button = driver.find_element(By.CLASS_NAME, 'pega-login')
        login_button.click()
        time.sleep(2)

        # time config
        _end_time = config.time.end_time
        _start_time = config.time.start_time
        if not _end_time:
            end_time = datetime.now().replace(hour=8, minute=0, second=0, microsecond=0)
        else:
            end_time = datetime.strptime(_end_time, date_time_format)
        if not _start_time:
            start_time = end_time - timedelta(days=1)
        else:
            start_time = datetime.strptime(_start_time, date_time_format)

        # Iterate over each domain
        download_directory = config.path.download  # Update the download directory path
        for desired_url in domains:
            try:
                # Open the desired URL and generate the report
                driver.get(desired_url)
                time.sleep(3)
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(
                    (By.CSS_SELECTOR, "li.dropdown>a.dropdown-toggle.querydate-menu"))).click()
                time.sleep(3)
                WebDriverWait(driver, 20).until(
                    EC.visibility_of_element_located((By.CSS_SELECTOR, "li>a#query_time"))).click()

                # Select the start date
                start_date_input = driver.find_element(By.CSS_SELECTOR,
                                                       "div#datepicker_begin>input.form-control[title='Select manually']")
                driver.execute_script("arguments[0].removeAttribute('readonly')", WebDriverWait(driver, 3).until(
                    EC.element_to_be_clickable(
                        (By.CSS_SELECTOR, "div#datepicker_begin>input.form-control[title='Select manually']"))))
                start_date_input.clear()
                start_date_input.send_keys(start_time.strftime(date_time_format))
                time.sleep(1)

                # Select the end date
                driver.find_element(By.ID, "chk_nowdate").click()
                end_date_input = driver.find_element(By.CSS_SELECTOR,
                                                     "div#datepicker_end>input.form-control[title='Select manually']")
                driver.execute_script("arguments[0].removeAttribute('readonly')", WebDriverWait(driver, 3).until(
                    EC.element_to_be_clickable(
                        (By.CSS_SELECTOR, "div#datepicker_end>input.form-control[title='Select manually']"))))
                end_date_input.clear()
                end_date_input.send_keys(end_time.strftime(date_time_format))

                # Click the report button
                driver.find_element(By.ID, 'submit_bnt').click()
                time.sleep(1)

                WebDriverWait(driver, 20).until(EC.invisibility_of_element((By.CSS_SELECTOR, "th>div.loading-font")))
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(
                    (By.CSS_SELECTOR, "ul.nav.navbar-nav.navbar-right>li.export_lineviewreport")))
                driver.find_element(By.CLASS_NAME, "export_lineviewreport").click()
                time.sleep(2)

                # Append the 'Downloads' folder to the home directory
                # After clicking the export button, wait for the file to be downloaded
                start_time_real = time.time()
                while time.time() - start_time_real < max_download_wait_time:
                    excel_files = [f for f in os.listdir(download_directory) if f.endswith('.xlsx')]
                    if excel_files:
                        time.sleep(2)  # Wait a bit to ensure the file is completely downloaded
                        latest_file = max([os.path.join(download_directory, f) for f in excel_files],
                                          key=os.path.getmtime)
                        shutil.move(latest_file, os.path.join(output_folder, os.path.basename(latest_file)))
                        print(f"Moved file: {os.path.basename(latest_file)}")
                        break
                    time.sleep(1)
                else:
                    print(f"Timeout: No Excel file was downloaded for URL {desired_url}")
            except Exception as e:
                print(f"Error processing URL {desired_url}: {str(e)}")
            # Close the browser

        # Read the downloaded Excel files and append their content to 'summary.xlsx'
        try:
            combined_wb = Workbook()
            combined_ws = combined_wb.active
            current_row = 1
            for file in os.listdir(output_folder):
                if file.endswith(".xlsx"):
                    file_path = os.path.join(output_folder, file)
                    wb = load_workbook(file_path)
                    ws = wb.active
                    for merged_range in ws.merged_cells.ranges:
                        if merged_range.min_row <= 4:
                            combined_ws.merge_cells(
                                start_row=current_row + merged_range.min_row - 1,
                                start_column=merged_range.min_col,
                                end_row=current_row + merged_range.max_row - 1,
                                end_column=merged_range.max_col
                            )
                    for i, row in enumerate(ws.iter_rows(), start=1):
                        for j, cell in enumerate(row, start=1):
                            new_cell = combined_ws.cell(row=current_row, column=j, value=cell.value)
                            # new_row=[]
                            # for cell in row:
                            # new_cell=combined_ws.cell(row=combined_ws.max_row+1, column=cell.column, value=cell.value)
                            if cell.has_style:
                                new_cell.font = copy(cell.font)
                                new_cell.border = copy(cell.border)
                                new_cell.fill = copy(cell.fill)
                                new_cell.number_format = copy(cell.number_format)
                                new_cell.protection = copy(cell.protection)
                                new_cell.alignment = copy(cell.alignment)
                            # new_row.append(new_cell)
                        # combined_ws.append( new_row)
                        current_row += 1
                    current_row += 1

            excel_file = os.path.join(output_folder,
                                      f"{start_time.strftime('%Y-%m-%d')}_{end_time.strftime('%Y-%m-%d')}_Detailed_YR_and_RR Report.xlsx")
            combined_wb.save(excel_file)

            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            data = []
            for row in range(1, ws.max_row):
                cell_b = ws.cell(row=row, column=2)
                cell_c = ws.cell(row=row, column=3)

                if cell_b.value == "Yield Rate" and cell_c.value == "Retest Rate(Max.)":
                    project = ws.cell(row=row - 3, column=1).value
                    yr_value = ws.cell(row=row + 1, column=2).value
                    rr_value = ws.cell(row=row + 1, column=3).value
                    if yr_value is not None and rr_value is not None:
                        data.append((project, yr_value, rr_value))
            doc_tables = Document()
            if data:
                # Create a new document for the tables
                for project, yr_value, rr_value in data:
                    paragraph = doc_tables.add_paragraph()
                    run = paragraph.add_run(f"{project}")
                    run.bold = True

                    table = doc_tables.add_table(rows=2, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Yield Rate'
                    hdr_cells[1].text = 'Retest Rate(Max.)'

                    row_cells = table.rows[1].cells
                    row_cells[0].text = str(yr_value)
                    row_cells[1].text = str(rr_value)

                    doc_tables.add_paragraph()  # Add a blank line between tables
                    if float(rr_value.strip("%")) >= 3:
                        reason_paragraph = doc_tables.add_paragraph()
                        reason_run = reason_paragraph.add_run('The reason fail more than 3% is: ')
                        reason_run.bold = True
            # Save the document with tables
            doc_tables.save(os.path.join(output_folder, 'aaa.docx'))

            # Open the example document
            doc_example = Document('example.docx')
            # Format start_time and end_time
            start_formatted = start_time.strftime('%I:%M %p of %A, %B %d, %Y')
            end_formatted = end_time.strftime('%I:%M %p of %A, %B %d, %Y')
            # Create the new paragraph text
            new_paragraph_text = f"from {start_formatted} to {end_formatted}"
            # Find and replace the paragraph
            for paragraph in doc_example.paragraphs:
                if "from 8:00 am of Thursday, July 11, 2024 to 8:00 am of Friday, July 12, 2024" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "from 8:00 am of Thursday, July 11, 2024 to 8:00 am of Friday, July 12, 2024",
                        new_paragraph_text
                    )
            # Find the paragraph containing "List of production lines operating"
            target_paragraph = None
            for i, paragraph in enumerate(doc_example.paragraphs):
                if "List of production lines operating" in paragraph.text:
                    # Copy content from aaa.docx
                    doc_aaa = Document(os.path.join(output_folder, 'aaa.docx'))
                    for element in doc_aaa.element.body:
                        doc_example.element.body.insert(i + 1, element)
                        i += 1
                    break

            if target_paragraph:
                # Get the index of the target paragraph
                index = doc_example.paragraphs.index(target_paragraph)

                # Copy content from aaa.docx
                doc_aaa = Document(os.path.join(output_folder, 'aaa.docx'))
                for element in doc_aaa.element.body:
                    doc_example.element.body.insert(index + 1, element)
                    index += 1

            # Save the modified example document
            doc_example.save(os.path.join(output_folder, 'output.docx'))
            outlook = win32com.client.Dispatch("Outlook.Application")

            doc = Document(os.path.join(output_folder, 'output.docx'))
            html_content = "<html><body>"

            for element in doc.element.body:
                if element.tag.endswith('p'):
                    html_content += process_paragraph(element)
                elif element.tag.endswith('tbl'):
                    html_content += "<table border='1' style='border-collapse: collapse; background-color: #E6F3FF;'>"
                    for row in element.xpath('.//w:tr'):
                        html_content += "<tr>"
                        for cell in row.xpath('.//w:tc'):
                            html_content += "<td style='padding: 5px; color: #4682B4;'>"
                            for paragraph in cell.xpath('.//w:p'):
                                html_content += process_paragraph(paragraph)
                            html_content += "</td>"
                        html_content += "</tr>"
                    html_content += "</table>"

            html_content += "</body></html>"
            mail = outlook.CreateItem(0)
            mail.To = '; '.join([_item.strip() for _item in config.email.TO])
            mail.CC = '; '.join([_item.strip() for _item in config.email.CC])
            mail.Subject = "Detailed Yield Rate and Retest Rate Report"
            mail.HTMLBody = html_content
            mail.Attachments.Add(excel_file)
            mail.Display(True)
        except Exception as e:
            print(f"Error processing Excel files: {str(e)}")
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")
    finally:
        try:
            driver.quit()
        except:
            pass
        messagebox.showinfo("Program Finish", "The program has finished running.")
